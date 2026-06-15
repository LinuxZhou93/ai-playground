import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def add_heading_with_style(doc, text, level):
    h = doc.add_heading(text, level)
    # 简单调整下颜色，让它看起来专业点
    for run in h.runs:
        run.font.color.rgb = RGBColor(15, 76, 129) # 经典深蓝
    return h

def generate_curriculum_doc():
    doc = docx.Document()
    
    # Title
    t = doc.add_heading('《Carbon-X：石墨烯未来科学家》16课时细化课程开发框架', 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('提交人：周林 | 提交日期：2026年5月20日')
    
    # 1. 材料研究评估
    add_heading_with_style(doc, '一、 邮寄材料（蒙烯玻纤加热片）研究评估与教学转化结论', 1)
    doc.add_paragraph('针对前期邮寄的“蒙烯玻纤加热片”（220V，44W，升温速度1.5℃/s），我进行了深度的拆解与实测。结论如下：')
    
    p = doc.add_paragraph()
    p.add_run('1. 材料定性：').bold = True
    p.add_run('该材料升温极快、面状发热极其均匀，远超传统电阻丝，具有极强的视觉冲击力与“黑科技”属性。')
    p.add_run('完全可以且非常适合作为核心教具引入课程。')
    p.add_run('\n2. 教学转化痛点与【破局对策】：').bold = True
    p.add_run('\n痛点：220V直供且无温控模块，持续通电容易导致材料烧毁，且初中生实操存在安全隐患。\n')
    p.add_run('对策（反向教学法）：我们不规避这个缺陷，反而将其转化为最高阶的“工程挑战”。课程前期（L1-L8）通过安全隔离电源进行物理探究；课程后期（L13-L16）要求学生引入 Arduino/Micro:bit 主控与继电器，亲手为发热片设计“智能温控断电保护系统”。这样，课程就从简单的“材料展示”升级为了“软硬结合的自动化工程”。')

    # 2. 详细大纲
    add_heading_with_style(doc, '二、 16课时细化教学大纲 (对标顶级创客标准)', 1)
    
    # 模块 1
    add_heading_with_style(doc, '模块一：微观觉醒 —— 材料科学基础 (L1 - L4)', 2)
    doc.add_paragraph('目标：建立对石墨烯二维结构的认知，理解材料导电性差异。\n'
                      'L1《碳原子的魔术》：认识石墨烯六角蜂窝结构，对比石墨与金刚石。\n'
                      'L2《绝缘与导电的跨界》：探究玻纤（绝缘）与石墨烯（导电）的复合工艺原理。\n'
                      'L3《电极的秘密》：分析银浆/铜箔等导电介质对发热效率的影响。\n'
                      'L4《面状发热的真面目》：学生使用万用表，实测加热片不同坐标网格的电阻，绘制电阻云图。')

    # 模块 2
    add_heading_with_style(doc, '模块二：能量密码 —— 物理与热力学探究 (L5 - L8)', 2)
    doc.add_paragraph('目标：通过实测数据，验证焦耳定律，理解热传导与红外辐射。\n'
                      'L5《焦耳定律实战》：实测不同输入电压下的升温曲线，验证 Q=I²Rt。\n'
                      'L6《热力学视觉化》：结合红外热成像仪，观察“面状发热”的均匀热辐射扩散过程。\n'
                      'L7《能量的接力赛》：探究热传导率，测试石墨烯穿透不同介质（木板、亚克力、棉布）的热衰减。\n'
                      'L8《生命光波》：学习远红外波段特性，探讨其在医疗领域的应用基础。')

    # 模块 3
    add_heading_with_style(doc, '模块三：极端场景 —— 跨学科工程应用 (L9 - L12)', 2)
    doc.add_paragraph('目标：运用设计思维（Design Thinking），解决真实世界的极端问题。\n'
                      'L9《火星基地的寒夜》：为火星探索车（结合众擎机器人概念）设计石墨烯保温层。\n'
                      'L10《电车冬季卫士》：模拟新能源车电池低温掉电，设计电池包加热裹身。\n'
                      'L11《可穿戴医疗设备》：设计一款智能发热护膝，学习人体工程学与软性材料封装。\n'
                      'L12《零碳建筑供暖》：模拟隐形地暖系统，计算能效比与碳减排量。')

    # 模块 4
    add_heading_with_style(doc, '模块四：数字大脑 —— 智能控制系统闭环 (L13 - L16)', 2)
    doc.add_paragraph('目标：软硬结合，解决核心“温控”痛点，完成最终产品交付。\n'
                      'L13《环境的感知触角》：引入 NTC 热敏传感器，学习读取环境模拟量温度。\n'
                      'L14《逻辑驱动物理》：学习继电器（Relay）工作原理，通过代码控制发热片的通断电。\n'
                      'L15《PID 恒温算法初探》：编写算法，实现 50℃ 恒温控制，解决持续升温烧毁问题。\n'
                      'L16《Carbon-X 创客发布会》：学生展示最终带有温控系统的智能发热产品原型。')

    # 3. 研发分工建议
    add_heading_with_style(doc, '三、 研发协同与北京支持诉求', 1)
    doc.add_paragraph('基于刘总“扛起研发重任”的指示，我将全力推进此框架的落地。为确保课程达到最高专业度，现明确协同分工如下：')
    doc.add_paragraph('1. 我方（周林团队）主导：\n'
                      '   - 整体课程框架与教学法设计。\n'
                      '   - 模块三（工程场景）与模块四（智能控制）的软硬件结合研发。\n'
                      '   - 学生端实验单与课件排版设计。\n'
                      '2. 需北京方（方老师等）协助支持：\n'
                      '   - 模块一与模块二中，深度的材料科学理论验证，确保化学/物理原理不出现常识性错误。\n'
                      '   - 提供官方或实验室内部的石墨烯显微图片、制备工艺短视频等作为教学素材。\n'
                      '   - 如果北京有成熟的“火星机器人”课程内容，希望能同步过来，以便我将石墨烯的L9课程与机器人课程做联动串联。')

    desktop_path = os.path.expanduser('~/Desktop/Carbon-X_细化课程开发框架_12点提交版.docx')
    doc.save(desktop_path)
    return desktop_path

if __name__ == "__main__":
    path = generate_curriculum_doc()
    print(f"Document successfully created at: {path}")
